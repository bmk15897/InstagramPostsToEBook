from docx import Document
from docx.shared import Inches
import json
import regex as re
import time

from instagram_scraper import InstagramScraper
from instagram_scraper.constants import *

#mention the username here. If private, add login_user and login_pass.

userName = 'chihiro.whispers'
login_user = None
login_pass = None

args = {
	'usernames': [userName],
	'destination': userName,
	'login_user': login_user,
	'login_pass': login_pass,
	'quiet': True,
	'maximum': 0,
	'retain_username': False,
	'media_metadata': True,
	'media_types': ['image'],
	'latest': False,
	'profile_metadata': True
}

scraper = InstagramScraper(**args)
scraper.scrape()

with open(userName+"\\"+userName+".json", encoding="utf8") as f:
  data = json.load(f)

document = Document()

document.add_heading('@'+userName, 0)
document.add_heading(data["GraphProfileInfo"]["info"]["biography"], 1)

cnt = 0
for i in range(len(data["GraphImages"])):
	if data["GraphImages"][i]["__typename"] == "GraphImage":
		cnt+=1
		readable = time.ctime(data["GraphImages"][i]["taken_at_timestamp"])
		document.add_heading('Timestamp:	'+readable, level=3)
		tempData = data["GraphImages"][i]["display_url"].split('?')
		tempData = tempData[0].split('/')
		tempData = tempData[-1]
		document.add_picture(userName+"\\"+tempData, width=Inches(6))
		
		
		if len(data["GraphImages"][i]["edge_media_to_caption"]["edges"])>0:
			p = document.add_paragraph('')
			p.add_run(data["GraphImages"][i]["edge_media_to_caption"]["edges"][0]["node"]["text"].strip()).bold = True
		
		
		document.add_page_break()

document.add_heading('Id: '+data["GraphProfileInfo"]["info"]["id"], 0)
print(cnt)
document.save(userName+'.docx')